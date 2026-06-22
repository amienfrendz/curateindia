from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('Technical Specification Document', level=0)
doc.add_heading('Conversational Visual Search for CurateIndia', level=1)
doc.add_paragraph('')

meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate([('Author', 'CurateIndia Engineering'), ('Date', 'June 2026'), ('Status', 'Draft'), ('Version', '1.0')]):
    meta_table.rows[i].cells[0].text = k
    meta_table.rows[i].cells[1].text = v
doc.add_paragraph('')

# 1. Problem Statement
doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph(
    'Current travel search is keyword and filter-based. Users searching for curated stays have deeply personal, '
    'multi-dimensional criteria that cannot be expressed through dropdowns or simple text queries.')
doc.add_paragraph('Example user criteria:')
for c in [
    'Spatial: "Big gardens for kids to run", "isolated from neighbors", "end of cul-de-sac"',
    'Visual: "No carpets", "large airy rooms", "stairs with railings for elderly"',
    'Experiential: "Quiet - no street noise", "well-maintained vs stock photos"',
    'Neighborhood: "Old trees, cafes with character", "fringe of popular spots"',
    'Value: "Not too commercial/polished - the sweet spot between cheap and overpriced"',
]:
    doc.add_paragraph(c, style='List Bullet')
doc.add_paragraph(
    'These criteria require understanding visual content (photos), experiential content (reviews), '
    'and spatial context (maps/neighborhood data) - not just structured listing metadata. '
    'A static attribute schema cannot anticipate all possible query dimensions. The system must be dynamic.')

# 2. Solution Options
doc.add_heading('2. Solution Options Evaluated', level=1)

# Option A
doc.add_heading('2.1 Option A: Real-time Google Places Contextual Search (per query)', level=2)
doc.add_paragraph(
    'At search time, for each candidate property, call Google Places API with contextualContents field. '
    'Google returns AI-selected photos and review snippets relevant to the user\'s specific query. '
    'No pre-computation needed - Google does the multimodal reasoning on their side.')
doc.add_paragraph('Flow:')
for s in [
    'User query: "quiet place with big garden and no carpets"',
    'Our LLM extracts intent -> narrow to 15-20 candidate properties by location/vibe',
    'For each candidate: call Google Places API with contextualContents, passing the user query context',
    'Google returns relevant photos + review excerpts that match the query',
    'LLM reasons over the contextual evidence to rank and explain matches',
]:
    doc.add_paragraph(s, style='List Number')
doc.add_paragraph('Pros: Zero pre-computation, always fresh data, Google handles multimodal reasoning, scales to any property count.')
doc.add_paragraph('Cons: Per-query cost ($0.10/search at 20 properties), latency (15-20 sequential API calls), dependent on Google index quality, cannot leverage our editorial content.')

# Option B
doc.add_heading('2.2 Option B: CLIP + Review Embeddings (RAG approach)', level=2)
doc.add_paragraph(
    'Pre-embed all property photos with CLIP and all review sentences with a text embedding model. '
    'At query time, embed the user query and do nearest-neighbor retrieval across both photo and text vectors. '
    'LLM reasons over retrieved evidence.')
doc.add_paragraph('Flow:')
for s in [
    'One-time: CLIP-embed all gallery photos (880 vectors) + embed all review sentences (~8000 vectors)',
    'User query -> embed with CLIP (for visual) + text model (for experiential)',
    'Cosine similarity search -> top 30 matching photos + top 30 matching review sentences',
    'Group evidence by property -> send top 15 properties with evidence to LLM',
    'LLM reasons and ranks with explanations',
]:
    doc.add_paragraph(s, style='List Number')
doc.add_paragraph('Pros: Zero per-query API cost, sub-second retrieval, works offline, handles ANY query dimension dynamically, full control over data and ranking.')
doc.add_paragraph('Cons: Requires embedding computation (one-time), CLIP may miss subtle visual details, limited to cached photos/reviews (not live data).')

# Option C
doc.add_heading('2.3 Option C: Google Gemini Grounding with Maps', level=2)
doc.add_paragraph(
    'Use Google Gemini API with "Grounding with Google Maps" - ask Gemini any natural language question '
    'and it automatically pulls relevant Google Maps data (reviews, photos, business info, neighborhood context) to answer.')
doc.add_paragraph('Flow:')
for s in [
    'User query -> send directly to Gemini API with Maps grounding enabled',
    'Gemini internally searches Google Maps, reviews, photos for relevant evidence',
    'Returns grounded answer with citations to specific reviews/photos',
    'We filter/re-rank against our curated property list',
]:
    doc.add_paragraph(s, style='List Number')
doc.add_paragraph('Pros: Most powerful - access to all of Google Maps data, 300M+ reviews, all photos. Handles spatial queries natively. Minimal code to implement.')
doc.add_paragraph('Cons: Highest per-query cost (~$0.01-0.03), may return non-curated properties (needs filtering), less control over ranking/UX, currently available only in US/India.')

# Option D
doc.add_heading('2.4 Option D: Microsoft/Azure-first Approach', level=2)
doc.add_paragraph('Feasibility assessment of a Microsoft-ecosystem solution:')
doc.add_heading('Azure Maps (Bing Maps replacement):', level=3)
for s in [
    'POI Search: Search for places by name, category, location. Comparable to Google Places basic search.',
    'NO equivalent to Google contextualContents - no AI-powered photo/review analysis.',
    'NO equivalent to Grounding with Google Maps - cannot ask "is this on a quiet street?"',
    'Reviews/photos: Very limited corpus compared to Google. Not viable for travel visual search.',
    'Bing Places API: Only for managing YOUR OWN business listings - not for third-party property discovery.',
]:
    doc.add_paragraph(s, style='List Bullet')
doc.add_heading('Azure AI Services (what IS viable):', level=3)
for s in [
    'Azure OpenAI GPT-4o with vision: Can analyze property photos at query time. Same capability as our CLIP approach but pay-per-call.',
    'Azure AI Search: Enterprise vector search with hybrid (keyword + vector) retrieval. Replaces Pinecone/Qdrant.',
    'Azure OpenAI text-embedding-3-small: Text embeddings for reviews. $0.02/1M tokens.',
    'Overture Maps (Microsoft-backed open data): Base POI data, free. But no reviews/photos.',
]:
    doc.add_paragraph(s, style='List Bullet')
doc.add_paragraph(
    'Verdict: Microsoft/Azure CAN replicate the embeddings approach (Option B) with Azure AI Search + Azure OpenAI. '
    'However, it CANNOT replicate the spatial/neighborhood intelligence that Google provides via Maps Grounding. '
    'For a travel product, this is a critical gap. The neighborhood, noise, isolation, and "character" dimensions '
    'all require geographic context that only Google currently offers via API. '
    'Azure is viable for the visual + experiential layers, but you would need Google for spatial intelligence regardless.')

# Option E
doc.add_heading('2.5 Option E: Hybrid (Recommended for Prototype)', level=2)
doc.add_paragraph(
    'Combine embeddings (free, handles 70% of queries) with Google APIs (paid, handles spatial + verification). '
    'Cost-optimize by using free layers first, paid layers only when needed.')
doc.add_paragraph('Flow:')
for s in [
    'User query -> LLM classifies intent dimensions: visual / experiential / spatial / combined',
    'Visual + experiential intents: CLIP + text embedding retrieval (free, instant, <100ms)',
    'Spatial intents: Google Gemini with Maps Grounding (~$0.01/call, only ~30% of queries)',
    'Verification (optional): Google contextualContents on top 3-5 results (~$0.005/call)',
    'LLM assembles all evidence -> ranked results with explanations citing specific photos/reviews',
]:
    doc.add_paragraph(s, style='List Number')
doc.add_paragraph('This gives best quality-to-cost ratio: free for most queries, paid only for spatial/verification.')

# 3. Prototype Path
doc.add_heading('3. Prototype Path (88 Properties)', level=1)
doc.add_paragraph('Goal: Validate that multi-modal search quality resonates with real users. Ship within 1-2 weeks. Near-zero cost.')

doc.add_heading('3.1 Implementation Steps', level=2)
for title_t, desc in [
    ('Step 1: CLIP Embed Photos (Day 1)',
     'Run CLIP ViT-B/32 locally on 880 cached gallery photos. Output: 880 x 512-dim vectors stored as JSON (~1.8MB). '
     'Runtime: ~10 min on CPU. Each vector tagged with property_slug + photo_index + category. Cost: $0.'),
    ('Step 2: Acquire Rich Review Data (Day 1-2)',
     'Current state: Only 543 review texts cached (Google returns max 5 per property). '
     'This is ~1% of actual review volume (avg 860 reviews/property exist on Google). Too thin for reliable vector search.\n\n'
     'Solution - two complementary approaches:\n\n'
     'A) Fetch Google generativeSummary for each property: Google synthesizes ALL reviews into a comprehensive '
     '200-500 word AI-generated description per place. One API call per property. Covers most-mentioned dimensions '
     'across hundreds of reviews (noise, space, value, accessibility, neighborhood). '
     'Cost: 88 calls within free 10K/month tier = $0.\n\n'
     'B) Contextual review fetches: Call Google Places API with contextualContents using 5 targeted queries per property '
     '("noise and quiet", "accessibility stairs elderly", "garden outdoor space", "maintenance condition", '
     '"neighborhood area surroundings"). Each returns different review subsets Google deems relevant to that context. '
     'Gets ~25 diverse review snippets per property vs the default 5. '
     'Cost: 88 x 5 = 440 calls within free tier = $0.\n\n'
     'Combined output per property: 1 comprehensive AI summary + 25 contextually-diverse review snippets. '
     'This gives far richer experiential signal than the current 5 generic reviews.'),
    ('Step 3: Embed Reviews (Day 2)',
     'Embed all review data using sentence-transformers all-MiniLM-L6-v2 (384-dim, runs locally on CPU):\n'
     '- Google generativeSummaries: split into sentences, embed each (~5-10 sentences per property = 440-880 vectors)\n'
     '- Contextual review snippets: embed each snippet (~25 per property = 2,200 vectors)\n'
     '- Total: ~2,600-3,100 vectors stored as JSON (~4MB)\n'
     '- Runtime: ~3-5 min locally. Cost: $0.\n\n'
     'Vector search at this scale: Pure NumPy in-memory cosine similarity. '
     'No database needed. Search time: <5ms for brute-force over 3,000 vectors. '
     'A vector DB (Pinecone, Qdrant) is unnecessary until 500K+ vectors (~5,000+ properties). '
     'Total infrastructure cost for review vector search: $0.'),
    ('Step 4: Search API Integration (Day 3-4)',
     'New retrieval pipeline in /api/search: (1) Embed user query with both CLIP and text model. '
     '(2) Cosine similarity against photo vectors -> top 30 photos across all properties. '
     '(3) Cosine similarity against review vectors -> top 30 sentences. '
     '(4) Score properties by evidence density. (5) Send top 15 to LLM with evidence context. '
     'Replace current keyword pre-filter. Keep LLM reasoning stage unchanged.'),
    ('Step 5: Spatial Layer (Day 5, optional)',
     'When LLM detects spatial intent in query ("neighborhood", "road", "isolated", "nearby"), '
     'call Google Gemini API with Maps Grounding for the candidate properties. Inject spatial context into LLM reasoning. '
     'Graceful fallback if API unavailable or budget exceeded.'),
    ('Step 6: Verification Layer (Day 5-6, optional)',
     'For final top 5 results, call Google Places contextualContents with the original user query. '
     'Google returns the most relevant photos and review snippets for that specific query. '
     'LLM uses this as confirmation/reranking signal. Only triggered if confidence is low.'),
]:
    doc.add_heading(title_t, level=3)
    doc.add_paragraph(desc)

doc.add_heading('3.2 Prototype Cost Breakdown', level=2)
t = doc.add_table(rows=9, cols=4)
t.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Item', 'One-time', 'Per Search', 'Monthly (1K searches)'),
    ('CLIP embeddings (local Python)', '$0', '-', '$0'),
    ('Review embeddings (local Python)', '$0', '-', '$0'),
    ('Vector search (in-memory NumPy)', '-', '<1ms', '$0'),
    ('LLM reasoning (GitHub Models free)', '-', '1 call', '$0 (150/day cap)'),
    ('Google Maps Grounding (~30% of queries)', '-', '$0.01', '~$3'),
    ('Google contextualContents (top 5, ~20% queries)', '-', '$0.025', '~$5'),
    ('Vercel hosting', '-', '-', '$0'),
    ('TOTAL', '$0', '$0-0.035', '$0-8/month'),
]):
    for j, val in enumerate(row):
        t.rows[i].cells[j].text = val

doc.add_heading('3.3 Alternative Prototype: Google Real-time Only (No Embeddings)', level=2)
doc.add_paragraph(
    'Simplest possible prototype - no embeddings, just call Google for every query. '
    'Trade cost for implementation speed.')
t2 = doc.add_table(rows=6, cols=3)
t2.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Item', 'Per Search', 'Monthly (1K searches)'),
    ('Google Text Search (find candidates)', '$0.005', '$5'),
    ('Google contextualContents x 20 properties', '$0.10', '$100'),
    ('LLM reasoning', '$0 (free tier)', '$0'),
    ('10K free requests/month covers', '-', 'First ~500 searches free'),
    ('TOTAL', '~$0.10', '$50-105'),
]):
    for j, val in enumerate(row):
        t2.rows[i].cells[j].text = val
doc.add_paragraph(
    'This is simpler to build (2-3 days vs 1 week) but 10x more expensive per search. '
    'Good for quick validation if you expect <500 searches/month (stays within free tier).')

# 4. Productization Path
doc.add_heading('4. Productization Path (500+ Properties, 10K+ searches/month)', level=1)

doc.add_heading('4.1 Architecture Changes from Prototype', level=2)
t3 = doc.add_table(rows=8, cols=3)
t3.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Component', 'Prototype (88 properties)', 'Production (500+)'),
    ('Vector storage', 'JSON file in-memory (~10KB)', 'Vector DB - Pinecone/Qdrant (free tier: 100K vectors)'),
    ('Photo source', 'Google Places cached locally', 'Host uploads + guest photos + Google Places'),
    ('Embedding compute', 'Local CPU one-time run', 'Cloud pipeline triggered on new content'),
    ('Search latency', '~300-500ms (NumPy + LLM)', '<200ms retrieval + streaming LLM'),
    ('LLM reasoning', 'GitHub Models free (150/day)', 'OpenAI/Azure OpenAI API ($0.003-0.005/query)'),
    ('Spatial intelligence', 'Google Gemini on-demand', 'Pre-computed neighborhood profiles + on-demand'),
    ('Quality monitoring', 'Manual testing', 'Automated quality metrics, A/B testing'),
]):
    for j, val in enumerate(row):
        t3.rows[i].cells[j].text = val

doc.add_heading('4.2 Production Ingestion Pipeline', level=2)
doc.add_paragraph('Runs automatically when new content arrives:')
for s in [
    'New photo uploaded -> CLIP embed -> upsert to vector DB with metadata (property_id, category, source, timestamp)',
    'New review received -> sentence-split -> embed -> upsert to vector DB with metadata',
    'Monthly: Re-fetch Google Places photos/reviews for all properties to catch new content',
    'Neighborhood refresh: Google Maps Grounding -> cache structured spatial profile per property',
    'Quality scoring: Flag low-quality or irrelevant photos for manual review',
]:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('4.3 Production Cost Estimate', level=2)
t4 = doc.add_table(rows=10, cols=4)
t4.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Item', 'Setup Cost', 'Monthly (10K searches)', 'Notes'),
    ('Vector DB (Pinecone free/starter)', '$0', '$0-70', 'Free up to 100K vectors; 500 props x 25 = 12.5K'),
    ('CLIP embedding (new photos)', '$5', '$2-5', 'GPU batch on new uploads only'),
    ('Text embedding (new reviews)', '$2', '$1-2', 'Incremental as reviews arrive'),
    ('LLM reasoning (GPT-4o-mini)', '$0', '$30-50', '10K queries x ~3K tokens avg'),
    ('Google Maps Grounding (30% queries)', '$0', '$30', '3K spatial queries x $0.01'),
    ('Google contextualContents (optional)', '$0', '$0-50', 'Only for low-confidence results'),
    ('Hosting (Vercel Pro)', '$0', '$20', 'Current plan'),
    ('Ingestion pipeline', '$0', '$5-10', 'Cloud functions on new content'),
    ('TOTAL', '~$7', '$90-210/month', ''),
]):
    for j, val in enumerate(row):
        t4.rows[i].cells[j].text = val

doc.add_heading('4.4 Microsoft Azure Alternative (Production)', level=2)
t5 = doc.add_table(rows=8, cols=3)
t5.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Component', 'Azure Service', 'Monthly Cost'),
    ('Vector search', 'Azure AI Search (Basic tier)', '$70-150'),
    ('Photo embeddings', 'Azure OpenAI GPT-4o vision OR local CLIP', '$5-10'),
    ('Text embeddings', 'Azure OpenAI text-embedding-3-small', '$2-5'),
    ('LLM reasoning', 'Azure OpenAI GPT-4o-mini', '$30-50'),
    ('Spatial intelligence', 'Azure Maps (basic POI only) + manual enrichment', '$10-20 + engineering time'),
    ('Hosting', 'Azure App Service OR keep Vercel', '$20-50'),
    ('TOTAL (no spatial equivalent)', '', '$140-285/month'),
]):
    for j, val in enumerate(row):
        t5.rows[i].cells[j].text = val
doc.add_paragraph(
    'Key finding: Azure is 40-60% more expensive than the Google hybrid approach AND lacks spatial intelligence. '
    'Microsoft has no equivalent to Google Maps Grounding or contextualContents. '
    'The only scenario where Azure wins: if you already have Azure enterprise commitments with credits, '
    'or if you need to keep all data within Microsoft ecosystem for compliance reasons.')
doc.add_paragraph(
    'Recommendation: Use Google for spatial/places intelligence (irreplaceable data moat), '
    'but the embeddings layer (CLIP, vector search) is provider-agnostic and can run anywhere.')

# 5. Search Quality
doc.add_heading('5. Expected Search Quality by Query Type', level=1)
t6 = doc.add_table(rows=10, cols=4)
t6.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('User Query', 'Primary Data Source', 'Expected Quality', 'Notes'),
    ('"Big gardens, space for kids"', 'Photos (CLIP)', 'High', 'CLIP excels at spatial/scene understanding'),
    ('"No carpets, hardwood floors"', 'Photos (CLIP)', 'High', 'Material/texture recognition is strong'),
    ('"Quiet, no street noise"', 'Reviews + Maps Grounding', 'High', 'Reviews mention noise; Maps shows road type'),
    ('"Isolated, far from neighbors"', 'Photos + Maps', 'Medium-High', 'Exterior photos show setting; Maps confirms'),
    ('"Nice neighborhood, old trees, cafes"', 'Google Maps Grounding', 'High', 'Google has rich neighborhood data'),
    ('"No internal stairs (elderly access)"', 'Photos + Reviews', 'Medium', 'Visible in photos; reviews sometimes mention'),
    ('"Well-maintained, not shabby"', 'Guest photos + Reviews', 'Medium-High', 'Recent photos vs stock; reviews cite condition'),
    ('"Sweet spot - not too commercial"', 'Reviews + Photo aesthetic', 'Medium', 'Sentiment + visual style tells a lot'),
    ('"View from property"', 'Photos (CLIP)', 'Medium', 'Only answerable if view was photographed'),
]):
    for j, val in enumerate(row):
        t6.rows[i].cells[j].text = val

# 6. Risks
doc.add_heading('6. Risks and Mitigations', level=1)
t7 = doc.add_table(rows=6, cols=3)
t7.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Risk', 'Impact', 'Mitigation'),
    ('Google Places photos include nearby landmarks, not the property', 'False visual matches', 'AI curation already done; host uploads in production'),
    ('CLIP misses subtle details (railings, carpet texture)', 'Incomplete matching for specific queries', 'LLM vision verification on top candidates (+$0.01/search)'),
    ('Some properties have <5 reviews', 'Low experiential signal for those properties', 'Weight photo evidence; structured host Q&A in production'),
    ('Multi-source retrieval + LLM adds latency', 'Slow UX (>3s)', 'Parallel retrieval; stream LLM response; cache patterns'),
    ('Google API cost spikes at scale', 'Budget overrun', 'Free embedding layer handles 70%; Google only for spatial + verification'),
]):
    for j, val in enumerate(row):
        t7.rows[i].cells[j].text = val

# 7. Timeline
doc.add_heading('7. Recommended Timeline', level=1)
t8 = doc.add_table(rows=7, cols=3)
t8.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Phase', 'Duration', 'Deliverable'),
    ('Prototype: CLIP + review embeddings', '3-4 days', 'Embedding pipeline, search API integration'),
    ('Prototype: Google spatial layer', '2 days', 'Maps Grounding for neighborhood queries'),
    ('Prototype: End-to-end testing', '2-3 days', '30+ diverse test queries benchmarked'),
    ('User validation', '2 weeks', '10-20 real travellers test, feedback collected'),
    ('Go/No-go decision', '-', 'Based on quality feedback + engagement metrics'),
    ('Productize (if go)', '4-6 weeks', 'Vector DB, ingestion pipeline, host uploads, monitoring'),
]):
    for j, val in enumerate(row):
        t8.rows[i].cells[j].text = val

# 8. Decision Points
doc.add_heading('8. Open Decision Points', level=1)
for d in [
    'Should results explain WHY they matched? ("Hardwood floors visible in photo 3, 4 reviews mention quiet surroundings") - adds user trust but increases LLM token cost.',
    'At what search volume does free tier break? ~150 LLM calls/day with GitHub Models. Above that, switch to paid ($30-50/month).',
    'Should hosts upload their own photos? Richer data but adds curation burden and moderation needs.',
    'Google real-time approach vs embeddings: Trade implementation simplicity for 10x higher per-query cost?',
    'Should spatial queries be real-time (fresh, $0.01/call) or pre-computed (stale, $0)? Depends on how often neighborhoods change.',
    'Provider strategy: Google-first (best data for travel) vs Azure-first (enterprise integration) vs agnostic (max flexibility)?',
]:
    doc.add_paragraph(d, style='List Number')

# 9. Appendix
doc.add_heading('9. Appendix: Cost Comparison Summary', level=1)
t9 = doc.add_table(rows=6, cols=5)
t9.style = 'Medium Shading 1 Accent 1'
for i, row in enumerate([
    ('Approach', 'Setup', 'Per Search', '1K searches/mo', '10K searches/mo'),
    ('A: Google real-time only', '$0', '$0.10', '$50-105', '$500-1050'),
    ('B: Embeddings only (no spatial)', '$0', '$0', '$0', '$0 (free tier)'),
    ('C: Gemini Grounding only', '$0', '$0.01-0.03', '$10-30', '$100-300'),
    ('E: Hybrid (recommended)', '$0', '$0-0.035', '$0-8', '$90-210'),
    ('D: Azure-first (no spatial)', '$0', '$0.003-0.01', '$3-10', '$140-285'),
]):
    for j, val in enumerate(row):
        t9.rows[i].cells[j].text = val

doc.save('scripts/TSD-Conversational-Visual-Search.docx')
print('Done: scripts/TSD-Conversational-Visual-Search.docx')
